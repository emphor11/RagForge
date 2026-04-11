import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class ExportService:
    def __init__(self, output_dir: str = "exports"):
        self.output_dir = output_dir
        # No longer creating local directories for the cloud version

    def _set_cell_background(self, cell, fill_color):
        """Helper to set cell background color using XML."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)

    def _add_page_number(self, paragraph):
        """Helper to add dynamic 'Page X of Y' to a paragraph."""
        run = paragraph.add_run()
        
        # XML for Page X
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText1 = OxmlElement('w:instrText')
        instrText1.set(qn('xml:space'), 'preserve')
        instrText1.text = "PAGE"
        run._r.append(instrText1)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # " of "
        paragraph.add_run(" of ")
        
        # XML for NumPages Y
        run2 = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = "NUMPAGES"
        run2._r.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar4)

    def generate_report(self, document_id: str, insight_data: dict) -> str:
        """
        Generates a professionally formatted DOCX report from RAG output.
        """
        doc = Document()

        # Document styles
        style = doc.styles['Normal']
        if hasattr(style, 'font'):
            style.font.name = 'Arial'
            style.font.size = Pt(10)

        profile = insight_data.get('contract_profile', {})
        document_type = profile.get('document_type', 'Contract')
        parties = profile.get('parties', [])
        effective_date = profile.get('effective_date', 'Not Detected')

        # Proper Document Title (Point 2)
        if len(parties) >= 2:
            doc_title = (
                f"{document_type} — {parties[0]} and {parties[1]}, "
                f"dated {effective_date}"
            )
        elif len(parties) == 1:
            doc_title = f"{document_type} — {parties[0]}, dated {effective_date}"
        else:
            doc_title = f"{document_type} — dated {effective_date}"

        # ---------------------------
        # FIRM HEADER BLOCK (Point 1)
        # ---------------------------
        header_title = doc.add_heading('CONTRACT REVIEW MEMORANDUM', 0)
        header_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Table Grid'

        # Populate Metadata Table
        rows = [
            ("Prepared by", "RAGForge Legal Intelligence"),
            ("Review date", datetime.now().strftime('%d %B %Y')),
            ("Document reviewed", doc_title),
            ("Prepared for", "Client (Confidential)")
        ]

        for i, (label, value) in enumerate(rows):
            meta_table.rows[i].cells[0].text = label
            meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            meta_table.rows[i].cells[1].text = value

        doc.add_paragraph("\n")

        # ---------------------------
        # EXECUTIVE SUMMARY (Point 4)
        # ---------------------------
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        formal_summary = (
            insight_data.get('formal_executive_summary') or
            insight_data.get('summary', 'Summary not available.')
        )
        doc.add_paragraph(formal_summary)

        # ---------------------------
        # CLAUSE SCORECARD (Point 3)
        # ---------------------------
        scorecard = insight_data.get('clause_scorecard', [])
        if scorecard:
            doc.add_heading('CLAUSE OVERVIEW & SCORECARD', level=1)
            score_table = doc.add_table(rows=1, cols=3)
            score_table.style = 'Table Grid'
            hdr_cells = score_table.rows[0].cells
            hdr_cells[0].text = 'Clause Type'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Risk Level'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True

            for item in scorecard:
                row = score_table.add_row().cells
                row[0].text = item.get('clause_type', 'N/A')
                status = item.get('status', 'Missing')
                row[1].text = status
                row[2].text = item.get('risk_level', 'None')

                # Apply background colors (Point 3)
                if status == 'Present':
                    self._set_cell_background(row[1], "C6EFCE")  # Light Green
                elif status == 'Partial':
                    self._set_cell_background(row[1], "FFEB9C")  # Light Amber
                else:
                    self._set_cell_background(row[1], "FFC7CE")  # Light Red

        doc.add_page_break()

        # ---------------------------
        # DETAILED FINDINGS (Point 5)
        # ---------------------------
        findings = insight_data.get('review_findings', [])
        if findings:
            doc.add_heading('DETAILED AUDIT FINDINGS', level=1)

            for f in findings:
                # Severity-based styling (Point 5)
                sev = str(f.get('severity', 'low')).lower()
                title_p = doc.add_paragraph()
                title_run = title_p.add_run(
                    f"[{sev.upper()}] {f.get('title', 'Finding')}"
                )
                title_run.bold = True

                if sev == 'high':
                    # Deep Red
                    title_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                elif sev == 'medium':
                    # Dark Amber
                    title_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

                # Explanation
                doc.add_paragraph(f.get('explanation', ''))

                # Source Quotes (Indented, italic, grey)
                quotes = f.get('source_quotes', [])
                if quotes:
                    for quote in quotes:
                        q_p = doc.add_paragraph()
                        q_p.paragraph_format.left_indent = Inches(0.3)
                        q_run = q_p.add_run(f'"{quote}"')
                        q_run.italic = True
                        q_run.font.color.rgb = RGBColor(128, 128, 128)

                # Border (thin horizontal line)
                hr_p = doc.add_paragraph("_" * 70)
                hr_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---------------------------
        # FOOTER & DISCLAIMER (Point 6)
        # ---------------------------
        section = doc.sections[0]
        footer = section.footer

        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Disclaimer (italics, grey)
        disclaimer_text = (
            "Disclaimer: This memorandum is prepared using AI-assisted "
            "analysis and does not constitute legal advice. "
            "Findings require review by a qualified legal professional."
        )
        run = footer_p.add_run(disclaimer_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True

        # Page numbering (Point 6)
        page_p = footer.add_paragraph()
        page_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._add_page_number(page_p)
        page_p.runs[0].font.size = Pt(8)
        page_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # ---------------------------
        # Save mechanism - Modified for Cloud
        # ---------------------------
        import io
        from app.services.supabase_storage import SupabaseStorage
        
        target = io.BytesIO()
        doc.save(target)
        target.seek(0)
        
        storage = SupabaseStorage()
        safe_name = document_id.replace(' ', '_').replace('/', '_')
        remote_path = f"exports/{safe_name}.docx"
        
        storage.upload_bytes(target.getvalue(), remote_path)
        
        # Return the public URL for the download from Supabase
        return storage.get_public_url(remote_path)
