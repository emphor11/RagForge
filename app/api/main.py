import asyncio
import json
import os
import shutil
import time
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from app.db.insight_store import InsightStore
from app.db.database import SessionLocal
from app.services.export_service import ExportService
from app.services.job_store import ensure_db_tables, job_store, utcnow
from app.services.live_analysis import run_live_analysis
from app.services.supabase_storage import SupabaseStorage


STREAM_HEADERS = {
    "Cache-Control": "no-cache",
    "Connection": "keep-alive",
    "X-Accel-Buffering": "no",
}

ACTIVE_JOB_STATUSES = {"queued", "processing", "failed"}


def get_allowed_origins():
    origins = {
        "http://localhost:5173",
        "http://127.0.0.1:5173",
    }

    for env_name in ("FRONTEND_URL", "FRONTEND_URLS"):
        raw_value = os.getenv(env_name, "")
        if not raw_value:
            continue

        for origin in raw_value.split(","):
            cleaned = origin.strip().rstrip("/")
            if cleaned:
                origins.add(cleaned)

    return sorted(origins)


def log_analysis_event(document_id: str, stage: str, detail: Optional[str] = None):
    message = f"[analysis] document={document_id} stage={stage}"
    if detail:
        message += f" detail={detail}"
    print(message)


def serialize_job_payload(job: Optional[dict]):
    if not job:
        return None

    return {
        "job_id": job["job_id"],
        "document_id": job["document_id"],
        "filename": job["filename"],
        "status": job["status"],
        "stage": job.get("stage"),
        "error": job.get("error"),
        "created_at": job["created_at"],
        "updated_at": job["updated_at"],
        "started_at": job.get("started_at"),
        "completed_at": job.get("completed_at"),
    }


def list_documents_with_jobs():
    store = InsightStore()
    completed_docs = store.list_all()
    active_jobs = job_store.list_pending_documents()

    seen_ids = {doc["id"] for doc in completed_docs}
    merged = completed_docs[:]
    for job_doc in active_jobs:
        if job_doc["id"] not in seen_ids:
            merged.append(job_doc)

    return sorted(
        merged,
        key=lambda item: float(item.get("upload_date", 0)),
        reverse=True,
    )


def ensure_generation_ready(generator):
    if not getattr(generator, "api_key", None):
        raise HTTPException(
            status_code=503,
            detail="Generation is not configured. Set GROQ_API_KEY before using this endpoint.",
        )


def is_deep_verify_enabled():
    return os.getenv("ENABLE_DEEP_VERIFY", "").strip().lower() in {"1", "true", "yes", "on"}


def sse_payload(payload: dict) -> str:
    return f"data: {json.dumps(payload)}\n\n"


def persist_uploaded_source(temp_path: Path, filename: str) -> str:
    storage = SupabaseStorage()
    if storage.is_configured():
        remote_source_path = f"uploads/{int(time.time())}_{filename}"
        storage.upload_file(str(temp_path), remote_source_path)
        return f"supabase://{remote_source_path}"
    return str(temp_path)


def update_job_state(
    job_id: str,
    *,
    status: str,
    stage: str,
    error: Optional[str] = None,
    completed: bool = False,
):
    changes = {
        "status": status,
        "stage": stage,
        "error": error,
    }
    if status == "processing":
        changes["started_at"] = utcnow()
    if completed:
        changes["completed_at"] = utcnow()
    return job_store.update_job(job_id, **changes)


def build_query_docs(query: str, stored_contract: Optional[dict]) -> list[str]:
    from app.core.review.legal_query import select_contract_query_docs
    from app.core.retrieval.bm25_retriever import BM25Retriever

    if not stored_contract:
        return []

    contract_query_docs = []
    if stored_contract.get("contract_profile"):
        contract_query_docs = select_contract_query_docs(query, stored_contract)

    if contract_query_docs:
        return [
            clause["clause_text"]
            for clause in contract_query_docs
            if clause.get("clause_text")
        ]

    analysis_chunks = stored_contract.get("analysis_chunks") or []
    if analysis_chunks:
        ranked_docs = BM25Retriever(analysis_chunks).retrieve(query, k=5)
        return [doc["content"] for doc in ranked_docs if doc.get("content")]

    raw_text = stored_contract.get("raw_text") or ""
    if raw_text.strip():
        return [raw_text[:15000]]

    return []


def run_deep_verification(document_id: str):
    if not is_deep_verify_enabled():
        raise HTTPException(
            status_code=503,
            detail="Deep verification is disabled in the free-hosted deployment.",
        )

    from app.evaluation.evaluator import InsightEvaluator

    store = InsightStore()
    data = store.load(document_id)
    if not data:
        raise HTTPException(status_code=404, detail="Document not found")

    insights = data.get("insights") if "insights" in data else data
    raw_text = data.get("raw_text", "")
    if not raw_text.strip():
        raise HTTPException(
            status_code=400,
            detail="Raw document text is unavailable for deep verification.",
        )

    try:
        evaluator = InsightEvaluator()
        evaluation = evaluator.run(insights, raw_text)
        review_audit = evaluator.evaluate_legal_review(
            data.get("review_findings", []),
            data.get("clauses", []),
            data.get("contract_profile", {}),
        )
    except ImportError as exc:
        raise HTTPException(
            status_code=503,
            detail="Deep verification dependencies are not installed in the hosted runtime.",
        ) from exc

    data["evaluation"] = evaluation
    data["review_audit"] = review_audit
    data["verification_mode"] = "manual_deep_verify"
    store.save(document_id, data)

    return {
        "document_id": document_id,
        "evaluation": evaluation,
        "review_audit": review_audit,
        "mode": "manual_deep_verify",
    }


app = FastAPI()


@app.on_event("startup")
def startup():
    ensure_db_tables()


app.add_middleware(
    CORSMiddleware,
    allow_origins=get_allowed_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
def health_check():
    storage = SupabaseStorage()
    config_status = {
        "supabase_storage": "configured" if storage.is_configured() else "missing",
        "postgres_db": "configured" if SessionLocal is not None else "missing",
        "groq_api": "configured" if os.getenv("GROQ_API_KEY") else "missing",
        "deep_verify": "enabled" if is_deep_verify_enabled() else "disabled",
    }

    return {
        "status": "healthy",
        "service": "ragforge-api",
        "config": config_status,
    }


@app.get("/")
def root():
    return {"status": "ok", "version": "2.0", "message": "RAGForge v2 Engine"}


@app.post("/documents/analyse")
async def analyse_document(request: Request, file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Uploaded file must include a filename.")

    document_id = file.filename.strip()
    if not document_id:
        raise HTTPException(status_code=400, detail="Uploaded file must include a valid filename.")

    uploads_dir = Path("uploads")
    uploads_dir.mkdir(parents=True, exist_ok=True)
    temp_path = uploads_dir / f"{int(time.time())}_{file.filename}"

    async def event_stream():
        ensure_db_tables()
        job = None
        loop = asyncio.get_running_loop()
        progress_queue: asyncio.Queue[dict] = asyncio.Queue()

        def queue_event(payload: dict):
            loop.call_soon_threadsafe(progress_queue.put_nowait, payload)

        def progress_callback(stage: str, progress: int, detail: Optional[str] = None):
            log_analysis_event(document_id, stage, detail)
            if job:
                update_job_state(
                    job["job_id"],
                    status="processing",
                    stage=stage,
                    error=None,
                )
            queue_event(
                {
                    "document_id": document_id,
                    "job_id": job["job_id"] if job else None,
                    "status": "processing",
                    "stage": stage,
                    "progress": progress,
                    "detail": detail,
                }
            )

        analysis_task = None

        try:
            initial_payload = {
                "document_id": document_id,
                "status": "processing",
                "stage": "uploading_source",
                "progress": 5,
                "detail": "Receiving document upload",
            }
            log_analysis_event(document_id, "uploading_source", "Receiving document upload")
            yield sse_payload(initial_payload)

            with open(temp_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            source_path = persist_uploaded_source(temp_path, file.filename)
            job = job_store.create_job(
                document_id=document_id,
                filename=file.filename,
                source_path=source_path,
            )
            job = update_job_state(
                job["job_id"],
                status="processing",
                stage="uploading_source",
                error=None,
            )
            log_analysis_event(document_id, "uploading_source", "Source saved durably")
            yield sse_payload(
                {
                    "document_id": document_id,
                    "job_id": job["job_id"],
                    "status": "processing",
                    "stage": "uploading_source",
                    "progress": 10,
                    "detail": "Source saved to durable storage",
                }
            )

            def run_and_persist():
                try:
                    run_live_analysis(
                        document_id=document_id,
                        file_path=str(temp_path),
                        progress_callback=progress_callback,
                    )
                    update_job_state(
                        job["job_id"],
                        status="completed",
                        stage="completed",
                        error=None,
                        completed=True,
                    )
                    log_analysis_event(document_id, "completed", "Analysis complete")
                    queue_event(
                        {
                            "document_id": document_id,
                            "job_id": job["job_id"],
                            "status": "completed",
                            "stage": "completed",
                            "progress": 100,
                            "detail": "Analysis complete",
                        }
                    )
                except Exception as exc:
                    message = str(exc) or "Document analysis failed."
                    update_job_state(
                        job["job_id"],
                        status="failed",
                        stage="failed",
                        error=message,
                    )
                    log_analysis_event(document_id, "failed", message)
                    queue_event(
                        {
                            "document_id": document_id,
                            "job_id": job["job_id"],
                            "status": "failed",
                            "stage": "failed",
                            "progress": 100,
                            "detail": message,
                            "error": message,
                        }
                    )

            analysis_task = asyncio.create_task(asyncio.to_thread(run_and_persist))

            while True:
                event = await progress_queue.get()
                if await request.is_disconnected():
                    break
                yield sse_payload(event)
                if event["stage"] in {"completed", "failed"}:
                    break

            if analysis_task:
                await analysis_task
        except Exception as exc:
            message = str(exc) or "Document analysis failed."
            if job:
                update_job_state(
                    job["job_id"],
                    status="failed",
                    stage="failed",
                    error=message,
                )
            log_analysis_event(document_id, "failed", message)
            yield sse_payload(
                {
                    "document_id": document_id,
                    "job_id": job["job_id"] if job else None,
                    "status": "failed",
                    "stage": "failed",
                    "progress": 100,
                    "detail": message,
                    "error": message,
                }
            )
        finally:
            await file.close()
            if temp_path.exists():
                temp_path.unlink()

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers=STREAM_HEADERS,
    )


@app.post("/upload", status_code=410)
def upload_route_retired():
    raise HTTPException(
        status_code=410,
        detail="Use POST /documents/analyse for streaming document analysis.",
    )


@app.get("/documents")
def list_documents():
    return list_documents_with_jobs()


@app.get("/jobs/{job_id}")
def get_job(job_id: str):
    job = job_store.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return serialize_job_payload(job)


@app.get("/documents/{document_id}/status")
def get_document_status(document_id: str):
    job = job_store.get_job_by_document(document_id)
    if job and job["status"] in ACTIVE_JOB_STATUSES:
        return {
            "document_id": document_id,
            "status": job["status"],
            "stage": job.get("stage"),
            "error": job.get("error"),
            "job": serialize_job_payload(job),
        }

    store = InsightStore()
    data = store.load(document_id)
    if data:
        return {
            "document_id": document_id,
            "status": "completed",
            "stage": "completed",
            "error": None,
            "job": serialize_job_payload(job) if job else None,
            "analysis_mode": data.get("analysis_mode", "groq_live"),
            "verification_mode": data.get("verification_mode"),
        }

    raise HTTPException(status_code=404, detail="Document or job not found")


@app.delete("/documents/{document_id}")
def delete_document(document_id: str):
    store = InsightStore()
    data = store.load(document_id)
    job = job_store.get_job_by_document(document_id)
    if not data and not job:
        raise HTTPException(status_code=404, detail="Document not found")

    if job and job["source_path"].startswith("supabase://"):
        storage = SupabaseStorage()
        if storage.is_configured():
            try:
                storage.supabase.storage.from_(storage.bucket_name).remove(
                    [job["source_path"].removeprefix("supabase://")]
                )
            except Exception as exc:
                print(f"❌ Failed to delete source upload for {document_id}: {exc}")
    elif job and os.path.exists(job["source_path"]):
        os.remove(job["source_path"])

    store.delete(document_id)
    job_store.delete_document_jobs(document_id)
    return {"message": "Document deleted successfully"}


@app.get("/insights/{document_id}")
def get_insights(document_id: str):
    store = InsightStore()
    data = store.load(document_id)

    if not data:
        job = job_store.get_job_by_document(document_id)
        if job and job["status"] in {"queued", "processing"}:
            raise HTTPException(
                status_code=202,
                detail="Document analysis is still in progress.",
            )
        if job and job["status"] == "failed":
            raise HTTPException(
                status_code=500,
                detail=job.get("error") or "Document analysis failed.",
            )
        raise HTTPException(status_code=404, detail="Not found")

    return data


@app.post("/documents/{document_id}/verify")
def verify_document(document_id: str):
    return run_deep_verification(document_id)


@app.get("/export/{document_id}")
def export_docx(document_id: str):
    store = InsightStore()
    data = store.load(document_id)
    if not data:
        raise HTTPException(status_code=404, detail="Document data not found")

    export_service = ExportService()
    try:
        export_url = export_service.generate_report(document_id, data)
        if not export_url:
            raise HTTPException(status_code=500, detail="Failed to generate export URL")
        return RedirectResponse(export_url)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/contracts/{document_id}/findings/{idx}/audit")
def update_finding_audit(document_id: str, idx: int, payload: dict):
    new_status = payload.get("status")
    reviewer_note = payload.get("reviewer_note")
    user_id = payload.get("user_id", "anonymous")

    store = InsightStore()
    if new_status:
        updated_finding = store.update_review_finding_status(
            document_id, idx, new_status, user_id
        )
    elif reviewer_note is not None:
        updated_finding = store.update_review_finding_note(
            document_id, idx, reviewer_note, user_id
        )
    else:
        raise HTTPException(status_code=400, detail="No status or note provided")

    if not updated_finding:
        raise HTTPException(status_code=404, detail="Finding not found")

    return {"message": "Audit logged successfully", "finding": updated_finding}


@app.get("/contracts/{document_id}/overview")
def get_contract_overview(document_id: str):
    data = InsightStore().load(document_id)

    if not data:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract_profile = data.get("contract_profile")
    if not contract_profile:
        raise HTTPException(
            status_code=404,
            detail="Contract profile not available for this document yet.",
        )

    return contract_profile


@app.get("/contracts/{document_id}/clauses")
def get_contract_clauses(document_id: str):
    data = InsightStore().load(document_id)

    if not data:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract_profile = data.get("contract_profile")
    if not contract_profile:
        raise HTTPException(
            status_code=404,
            detail="Contract profile not available for this document yet.",
        )

    clauses = data.get("clauses")
    if clauses is None:
        clauses = contract_profile.get("clause_index", [])

    return {
        "document_id": document_id,
        "clauses": clauses,
        "count": len(clauses),
    }


@app.get("/contracts/{document_id}/risks")
def get_contract_risks(document_id: str):
    data = InsightStore().load(document_id)

    if not data:
        raise HTTPException(status_code=404, detail="Contract not found")

    review_findings = data.get("review_findings", [])

    return {
        "document_id": document_id,
        "findings": review_findings,
        "count": len(review_findings),
    }


@app.get("/contracts/{document_id}/review-audit")
def get_contract_review_audit(document_id: str):
    data = InsightStore().load(document_id)

    if not data:
        raise HTTPException(status_code=404, detail="Contract not found")

    review_audit = data.get("review_audit")
    if not review_audit:
        raise HTTPException(
            status_code=404,
            detail="Contract review audit is not available for this document yet.",
        )

    return review_audit


@app.get("/reports")
def get_reports():
    store = InsightStore()
    docs = store.list_all()

    report_data = {
        "total_docs": len(docs),
        "total_risks": 0,
        "total_actions": 0,
        "risk_summary": {"high": 0, "medium": 0, "low": 0},
        "top_insights": [],
        "doc_list": [],
    }

    unique_insights = set()

    for doc in docs:
        full_data = store.load(doc["id"])
        if not full_data:
            continue

        details = full_data.get("insights") if "insights" in full_data else full_data

        risks = details.get("risks", [])
        report_data["total_risks"] += len(risks)
        for risk in risks:
            severity = risk.get("severity", "low").lower()
            if severity in report_data["risk_summary"]:
                report_data["risk_summary"][severity] += 1

        actions = details.get("recommended_actions", [])
        report_data["total_actions"] += len(actions)

        insights = details.get("key_insights", [])
        for insight in insights:
            insight_text = (
                insight.get("insight") if isinstance(insight, dict) else insight
            )
            if insight_text and insight_text not in unique_insights:
                unique_insights.add(insight_text)
                if len(report_data["top_insights"]) < 10:
                    report_data["top_insights"].append(insight_text)

        report_data["doc_list"].append(
            {
                "name": doc["id"],
                "summary": details.get("summary", "No summary available."),
            }
        )

    return report_data


@app.get("/analytics")
def get_analytics():
    from datetime import datetime, timedelta

    store = InsightStore()
    docs = store.list_all()

    analytics = {
        "docs_over_time": {},
        "risk_trend": {"last_week": 0, "this_week": 0, "direction": "stable"},
        "performance": {"avg_confidence": 0.0, "avg_response_time": 2.1},
        "usage": {"total_queries": 0, "total_analyses": len(docs)},
    }

    total_conf = 0
    now = datetime.now()
    one_week_ago = now - timedelta(days=7)

    for doc in docs:
        full_data = store.load(doc["id"])
        if not full_data:
            continue

        dt = datetime.fromtimestamp(doc["upload_date"])
        date_str = dt.strftime("%Y-%m-%d")
        analytics["docs_over_time"][date_str] = (
            analytics["docs_over_time"].get(date_str, 0) + 1
        )

        details = full_data.get("insights") if "insights" in full_data else full_data
        conf = (
            details.get("overall_confidence") or details.get("confidence_score") or 0.0
        )
        total_conf += conf

        risks_count = len(details.get("risks", []))
        if dt > one_week_ago:
            analytics["risk_trend"]["this_week"] += risks_count
        else:
            analytics["risk_trend"]["last_week"] += risks_count

    if len(docs) > 0:
        analytics["performance"]["avg_confidence"] = round(total_conf / len(docs), 2)

    if analytics["risk_trend"]["this_week"] < analytics["risk_trend"]["last_week"]:
        analytics["risk_trend"]["direction"] = "down"
    elif analytics["risk_trend"]["this_week"] > analytics["risk_trend"]["last_week"]:
        analytics["risk_trend"]["direction"] = "up"

    analytics["usage"]["total_queries"] = len(docs) * 4
    return analytics


class QueryRequest(BaseModel):
    query: str
    document_id: Optional[str] = None


class FindingStatusUpdate(BaseModel):
    status: str


class FindingNoteUpdate(BaseModel):
    reviewer_note: str


@app.post("/query")
def query_api(request: QueryRequest):
    from app.core.generation.structured_generator import StructuredGenerator

    query = request.query
    document_id = request.document_id
    stored_contract = InsightStore().load(document_id) if document_id else None

    if document_id and not stored_contract:
        raise HTTPException(status_code=404, detail="Document analysis not found.")

    docs = build_query_docs(query, stored_contract)
    if not docs:
        raise HTTPException(
            status_code=404,
            detail="No relevant context found for the query.",
        )

    generator = StructuredGenerator()
    ensure_generation_ready(generator)

    try:
        return generator.generate(query=query, docs=docs)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="Failed to generate a response for the query.",
        ) from exc


@app.patch("/contracts/{document_id}/findings/{finding_index}/status")
def update_contract_finding_status(
    document_id: str, finding_index: int, payload: FindingStatusUpdate
):
    allowed_statuses = {
        "open",
        "reviewed",
        "accepted",
        "dismissed",
        "escalated",
        "negotiate",
    }
    if payload.status not in allowed_statuses:
        raise HTTPException(status_code=400, detail="Invalid finding status.")

    store = InsightStore()
    updated_finding = store.update_review_finding_status(
        document_id, finding_index, payload.status
    )

    if not updated_finding:
        raise HTTPException(status_code=404, detail="Contract or finding not found.")

    return {
        "document_id": document_id,
        "finding_index": finding_index,
        "finding": updated_finding,
    }


@app.patch("/contracts/{document_id}/findings/{finding_index}/note")
def update_contract_finding_note(
    document_id: str, finding_index: int, payload: FindingNoteUpdate
):
    store = InsightStore()
    updated_finding = store.update_review_finding_note(
        document_id,
        finding_index,
        payload.reviewer_note.strip(),
    )

    if not updated_finding:
        raise HTTPException(status_code=404, detail="Contract or finding not found.")

    return {
        "document_id": document_id,
        "finding_index": finding_index,
        "finding": updated_finding,
    }


@app.get("/contracts/{document_id}/export")
def export_contract_report(document_id: str):
    data = InsightStore().load(document_id)
    if not data:
        raise HTTPException(status_code=404, detail="Contract not found")

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    import io

    doc = Document()

    title = doc.add_heading(f"Contract Review Report: {document_id}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    profile = data.get("contract_profile", {})
    doc.add_heading("Contract Overview", level=1)
    table = doc.add_table(rows=0, cols=2)

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value or "N/A")

    add_row(
        "Document Type", profile.get("document_type", "N/A").replace("_", " ").title()
    )
    add_row("Effective Date", profile.get("effective_date"))
    add_row("Governing Law", profile.get("governing_law"))
    add_row("Term Length", profile.get("term_length"))
    add_row("Parties", ", ".join(profile.get("parties", [])))

    findings = data.get("review_findings", [])
    if findings:
        doc.add_heading("Review Findings", level=1)
        for finding in findings:
            p = doc.add_paragraph()
            run = p.add_run(f"{finding.get('title', 'Untitled Finding')}")
            run.bold = True
            run.font.size = Pt(12)

            p = doc.add_paragraph()
            p.add_run(
                f"Type: {finding.get('finding_type', 'N/A').replace('_', ' ').title()} | "
            )
            sev = finding.get("severity", "N/A").upper()
            run_sev = p.add_run(f"Severity: {sev}")
            if sev == "HIGH":
                run_sev.font.color.rgb = RGBColor(200, 0, 0)

            doc.add_paragraph(finding.get("explanation", "No explanation provided."))

            if finding.get("source_quotes"):
                doc.add_paragraph("Source Quotes:")
                for quote in finding["source_quotes"]:
                    doc.add_paragraph(f'"{quote}"', style="List Bullet")

            doc.add_paragraph("-" * 30)

    clauses = data.get("clauses", [])
    if clauses:
        doc.add_heading("Clause Inventory", level=1)
        for clause in clauses:
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(
                f"{clause.get('title')} (Page {clause.get('page_number', 'N/A')})"
            )
            doc.add_paragraph(clause.get("clause_text", ""))

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)

    return StreamingResponse(
        target,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=RagForge_Report_{document_id.replace(' ', '_')}.docx"
        },
    )


frontend_path = os.path.join(os.getcwd(), "rag-ui", "dist")

if os.path.exists(frontend_path):
    app.mount(
        "/assets",
        StaticFiles(directory=os.path.join(frontend_path, "assets")),
        name="assets",
    )

    @app.get("/{rest_of_path:path}")
    async def react_app(rest_of_path: str):
        if rest_of_path.startswith(("api/", "docs", "openapi.json", "redoc")):
            raise HTTPException(status_code=404, detail="API route not found")

        index_file = os.path.join(frontend_path, "index.html")
        if os.path.exists(index_file):
            return FileResponse(index_file)
        return {
            "message": "Frontend build not found. Please run 'npm run build' in rag-ui."
        }

else:

    @app.get("/{rest_of_path:path}")
    async def build_missing(rest_of_path: str):
        return {
            "mode": "development",
            "message": "Backend is running, but UI build (dist) not found. Use 'npm run dev' for development.",
        }
